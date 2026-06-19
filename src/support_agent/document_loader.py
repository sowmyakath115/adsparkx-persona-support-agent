from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from support_agent.schemas import SourceLocation


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    location: SourceLocation


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

    def load_directory(self, directory: Path) -> list[LoadedDocument]:
        if not directory.exists():
            raise FileNotFoundError(f"Knowledge base directory not found: {directory}")

        documents: list[LoadedDocument] = []
        for path in sorted(directory.rglob("*")):
            if path.is_file() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                documents.extend(self.load_file(path, root=directory))
        if not documents:
            raise ValueError(f"No supported knowledge base documents found in {directory}")
        return documents

    def load_file(self, path: Path, root: Path | None = None) -> list[LoadedDocument]:
        suffix = path.suffix.lower()
        source = str(path.relative_to(root)) if root else path.name
        if suffix in {".md", ".txt"}:
            return self._load_text(path, source)
        if suffix == ".pdf":
            return self._load_pdf(path, source)
        if suffix == ".docx":
            return self._load_docx(path, source)
        raise ValueError(f"Unsupported file type: {path.suffix}")

    def _load_text(self, path: Path, source: str) -> list[LoadedDocument]:
        text = path.read_text(encoding="utf-8")
        sections = self._split_markdown_sections(text)
        if not sections:
            return [LoadedDocument(text=text, location=SourceLocation(source=source, section=path.stem))]
        return [
            LoadedDocument(text=section_text, location=SourceLocation(source=source, section=section_name))
            for section_name, section_text in sections
        ]

    @staticmethod
    def _split_markdown_sections(text: str) -> list[tuple[str, str]]:
        matches = list(re.finditer(r"^(#{1,3})\s+(.+)$", text, flags=re.MULTILINE))
        if not matches:
            return []
        sections: list[tuple[str, str]] = []
        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            section_name = match.group(2).strip()
            section_text = text[start:end].strip()
            if section_text:
                sections.append((section_name, section_text))
        return sections

    def _load_pdf(self, path: Path, source: str) -> list[LoadedDocument]:
        reader = PdfReader(str(path))
        loaded: list[LoadedDocument] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                loaded.append(LoadedDocument(text=text, location=SourceLocation(source=source, page=idx)))
        return loaded

    def _load_docx(self, path: Path, source: str) -> list[LoadedDocument]:
        doc = DocxDocument(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return [
            LoadedDocument(
                text="\n".join(paragraphs),
                location=SourceLocation(source=source, section=path.stem),
            )
        ]
